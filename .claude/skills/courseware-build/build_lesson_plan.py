#!/usr/bin/env python3
"""Generate the IoT Fundamental for Beginners Lesson Plan (LP) DOCX in the
Tertiary house format: cover page + Document Version Control Record + auto TOC
+ Arial 11pt body + colour-coded 2-day schedule tables (9:30am-6:30pm, 8
training hours/day, 1h lunch, tea within, final assessment Day 2 4:30pm).
Topics/labs come from course_data + data_labs so the LP stays aligned with the
deck, guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_labs import LABS
import prodoc
REPO=os.path.dirname(HERE); ASSETS=os.path.join(REPO,"courseware","assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in LABS if a['num'] in nums)

# ------------------------------------------------ 2-day schedule (single source of timing)
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","10:00",30,"admin","Welcome, trainer and learner introductions, learning outcomes, course outline, ground rules and mandatory digital attendance (AM)"),
    ("10:00","11:00",60,"topic","Topic 1 — Overview of Internet of Things (IoT): what is IoT, IoT devices, sensors and actuators, triggers and events (concepts + demo)"),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:00",105,"topic","Topic 1 (continued) — wireless communication technologies, IoT applications and use cases. Introduction to our IoT platform IoTFlow (iot.tertiaryinfotech.com): how it works — connect, visualise & control, automate with n8n"),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"lab","Topic 2 — Collect and Post Data to Cloud: cloud computing, IoT platforms, MQTT and REST API (concepts). Hands-on: "+lab_titles([1])),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","17:45",120,"lab","Hands-on: "+lab_titles([2])),
    ("17:45","18:30",45,"recap","Day 1 recap, Q&A and PM digital attendance"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30","9:45",15,"recap","Day 1 recap and mandatory digital attendance (AM)"),
    ("9:45","11:00",75,"lab","Topic 3 — Read Data and Remote Control from Cloud: REST API, MQTT subscribe, virtual pins (concepts). Hands-on: "+lab_titles([3])),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:00",105,"lab","Hands-on: "+lab_titles([4])+". Alerts and trigger rules"),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:00",60,"lab","Topic 4 — IoT Data Analytics and Visualization: dashboards, analytics, cybersecurity (concepts). Hands-on: "+lab_titles([5])),
    ("15:00","15:15",15,"break","Tea break"),
    ("15:15","16:15",60,"lab","Hands-on: "+lab_titles([6])),
    ("16:15","16:30",15,"assess","Course feedback and TRAQOM survey; Briefing for Assessment; Assessment digital attendance"),
    ("16:30","17:30",60,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book"),
    ("17:30","18:30",60,"assess","Practical Performance (PP) — hands-on IoT platform tasks, 1 hour, open book. End of class"),
 ]),
}

doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
    ("12",("1 March 2025"),"Previous release — ThingSpeak/ESP8266/Node-RED based courseware.",C.TRAINER),
    (C.VERSION.lstrip("v"),C.VERSION_DATE,
     "Revamped to the IoTFlow platform (iot.tertiaryinfotech.com): 6 labs covering device registration, "
     "MQTT/HTTP telemetry, virtual-pin remote control, dashboards and n8n + AI automation.",C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1): return doc.add_heading(text,level=level)

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Skills Framework TSC",f"{C.TSC_TITLE} ({C.TSC_CODE})"),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","2 days · 8 training hours per day (16 hours)"),
      ("Daily Timing","9:30 am – 6:30 pm (1-hour lunch; tea breaks within training time)"),
      ("Mode",f"Instructor-led, hands-on labs on the {C.PLATFORM} IoT platform ({C.PLATFORM_URL}) and n8n"),
      ("Trainer",C.TRAINER),
      ("Course Page",C.COURSE_URL)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Skills Framework — TSC Abilities and Knowledge",1)
doc.add_paragraph("This course is mapped to the Skills Framework TSC "+C.TSC_TITLE+" ("+C.TSC_CODE+"):")
for x in C.ABILITIES+C.KNOWLEDGE:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(x).font.size=Pt(10.5)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Final assessment is conducted on Day 2 from 4:30 pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None):
    cell.text=""; p=cell.paragraphs[0]
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),fill=fill)
        set_cell(cells[1],f"{mins} min",fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),fill=fill)
        if kind!="lunch": training+=mins
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours).")
    r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

H("Lab Reference (aligned to topics)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic","Abilities / Knowledge","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in LABS if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,fill=TOPIC_FILL)
    set_cell(cells[1],tp["tags"],fill=TOPIC_FILL)
    set_cell(cells[2],"; ".join(f"Lab {a['num']}: {a['title']}" for a in acts) if acts else "Concepts and demonstrations (no lab)")

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.FILE_STEM}.docx")
doc.save(OUT)
print("Saved",OUT)
