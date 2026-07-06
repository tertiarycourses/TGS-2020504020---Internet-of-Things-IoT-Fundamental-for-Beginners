#!/usr/bin/env python3
"""Generate the IoT Fundamental for Beginners Learner Guide as BOTH a Markdown
mirror (LEARNER-GUIDE.md at repo root) and a DOCX (courseware/LG-*.docx) from
one content stream, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
· Test it), plus platform setup, key IoT concepts and glossary. All content is
driven by course_data + data_labs, keeping the LG 100% aligned with the slide
deck, Lesson Plan and labs.
"""
import os, sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches as DocxInches

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_labs import LABS
import prodoc
REPO=os.path.dirname(os.path.dirname(os.path.dirname(HERE)))  # skill dir -> repo root
ASSETS=os.path.join(REPO,"courseware","assets")

def lab_slug(a):
    s=re.sub(r"[^a-z0-9]+","-",a["title"].lower()).strip("-")
    return f"lab-{a['num']:02d}-{s}.md"

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))
def img(fn,caption): B.append(("img",fn,caption))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  f"It provides step-by-step instructions for all 6 hands-on labs, organised by the four course topics. "
  f"Every lab runs on our IoT platform {C.PLATFORM} ({C.PLATFORM_URL}) and mirrors an official platform "
  f"tutorial ({C.TUTORIALS_URL}), so you can repeat the same steps any time after class.")
p("Use this guide alongside the course slides and the lab files in the labs/ folder of the course "
  "repository. No hardware is required — every lab can be completed with cURL or Python from any "
  "terminal — but an ESP32/ESP8266 or Raspberry Pi makes the experience real.")
p(f"Course page: {C.COURSE_URL}")

h1("Skills Framework")
p(f"This course is mapped to the Skills Framework TSC {C.TSC_TITLE} ({C.TSC_CODE}).")
h3("TSC Abilities")
bullets(C.ABILITIES)
h3("TSC Knowledge")
bullets(C.KNOWLEDGE)

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Key IoT Concepts at a Glance")
B.append(("dl",[
 ("Devices","Hardware — sensors, gadgets, appliances and machines — that collect and exchange data over the internet."),
 ("Sensors","The input side: they measure a physical quantity (temperature, humidity, motion, gas, light)."),
 ("Actuators","The output side: relays, motors, pumps, valves and LEDs that act on the physical world."),
 ("Triggers","Rules that turn readings into events — a threshold crossed or a device going offline fires an action."),
 ("n8n automation","A low-code workflow tool: device events fire flows that notify, log, call AI or control devices back."),
 ("Workflows","Chains of n8n nodes wired together — each node does one job (trigger, transform, notify, store, AI)."),
 ("AI","AI nodes inside workflows summarise readings, detect anomalies and recommend actions — analytics with zero code."),
 ("Dashboard for control","Web/mobile panels of widgets: display widgets show data; control widgets write virtual pins to command devices."),
]))

h1(f"Before You Start — {C.PLATFORM} Platform Setup")
h3("What you need")
bullets([
 f"A free account on {C.PLATFORM} — sign up at {C.PLATFORM_URL}.",
 "A modern browser. The dashboard is an installable PWA that also works on your phone.",
 "Optional hardware: ESP32 / ESP8266 / Raspberry Pi with a DHT temperature-humidity sensor.",
 "Access to an n8n instance for Lab 6 (cloud or self-hosted).",
])
h3("How the platform works")
bullets([st for st in C.HOW_IT_WORKS])

h3("Install the IoTFlow Python client (pip install iotflow)")
p("The official Python client connects any device that runs Python (Raspberry Pi, PC, Mac, Linux "
  "SBCs) to the platform. For microcontrollers (Arduino, ESP8266, ESP32) use the Arduino library — "
  "both speak the same protocol. Install it once before Lab 2:")
steps([
 ("Open a terminal (macOS/Linux: Terminal; Windows: Command Prompt or PowerShell) and check that Python 3 is installed.",
  "python --version    # or: python3 --version"),
 ("Install the client with MQTT support (recommended for this course).",
  'pip install "iotflow[mqtt]"'),
 ("Alternative — HTTP-only install (zero dependencies) if pip cannot reach paho-mqtt.",
  "pip install iotflow"),
 ("Verify the installation — the import must succeed silently.",
  'python -c "from iotflow import IoTFlow; print(\'iotflow OK\')"'),
])
p(f"Package on PyPI: {C.PYPI_URL}  ·  Source, README and examples: {C.PY_CLIENT_URL}")

h3("The telemetry_upload.py example script (used in Lab 2)")
p("This is the Python equivalent of the ESP8266_Telemetry_Upload Arduino sketch — it connects to "
  "the MQTT broker and publishes a reading every 10 seconds. Download it, fill in your broker "
  "host, device id and device token from the Add Device wizard, then run it:")
code(f"curl -O {C.PY_EXAMPLE_RAW}\npython telemetry_upload.py")
code('''import random
import time

from iotflow import IoTFlow

# ---- MQTT broker ----
# Use the course platform broker. (For a self-hosted broker on your LAN,
# use that machine's IP instead — find it with `ipconfig` / `ifconfig`.)
MQTT_BROKER = "iot.tertiaryinfotech.com"
MQTT_PORT = 1883

# ---- Device identity (from the "Add Device" wizard in IoTFlow) ----
DEVICE_ID = "test-device"
DEVICE_TOKEN = "dev_XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX"

PUBLISH_INTERVAL_S = 10

iot = IoTFlow(
    token=DEVICE_TOKEN,
    device_id=DEVICE_ID,
    mqtt_host=MQTT_BROKER,
    mqtt_port=MQTT_PORT,
)
iot.connect()  # MQTT connection runs in the background
print(f"Connected to MQTT broker {MQTT_BROKER}:{MQTT_PORT}")

while True:
    # TODO: replace with real sensor readings (DHT22, BME280, ADC, ...)
    temperature = round(28.5 + random.uniform(-1.0, 1.0), 1)
    humidity = round(65 + random.uniform(-5.0, 5.0), 1)
    voltage = 3.7

    iot.mqtt_publish(temperature=temperature, humidity=humidity, voltage=voltage)
    print(f"Published: temperature={temperature} humidity={humidity} voltage={voltage}")
    time.sleep(PUBLISH_INTERVAL_S)''')
note(f"Example script on GitHub: {C.PY_EXAMPLE_URL} — the Arduino original it mirrors: "
     f"{C.ARDUINO_EXAMPLE_URL}")

h3("Conventions used in every lab")
bullets([
 "Your device token (dev_...) is shown ONCE at registration — store it safely and never share it.",
 "Every API call authenticates with the header:  Authorization: Bearer dev_XXXXXXXXXXXX",
 "MQTT broker: iot.tertiaryinfotech.com, port 1883.  REST endpoint: /api/telemetry.",
 "Python code uses the official client: pip install \"iotflow[mqtt]\" — send()/virtual_write() over HTTP, connect()/mqtt_publish() over MQTT, @on_command for control.",
 f"Each lab mirrors a platform tutorial at {C.TUTORIALS_URL} — revisit them any time.",
])

# ---------------- per-topic, per-lab ----------------
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['tags']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    acts=[x for x in LABS if x["topic"]==t["num"]]
    if not acts:
        note("This topic is delivered through concepts, diagrams and trainer demonstrations — "
             "the hands-on labs begin in Topic 02.")
    for a in acts:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Learning outcome: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Uses: {a['services']}.)")
        if a.get("shot"):
            img(a["shot"], f"Platform walkthrough for Lab {a['num']} — {a['title']}")
        h3("Step-by-step")
        steps(list(a["steps"]))
        h3("Test it")
        p(a["test"])
        note(f"This lab mirrors the platform tutorial at {a['tutorial']} — "
             f"the full lab sheet is labs/{lab_slug(a)} in the course repository.")
        rule()

h1("After the Course")
bullets([
 f"Your {C.PLATFORM} account and projects remain yours — keep experimenting at {C.PLATFORM_URL}.",
 f"Revisit the tutorials at {C.TUTORIALS_URL} to repeat any lab.",
 "Connect a real ESP32 or Raspberry Pi at home using the Integrate page snippets.",
 "Extend your n8n flow: WhatsApp alerts, Google Sheets logging, AI agents.",
 "Assessment: WA (SAQ, 1 hour) + PP (practical, 1 hour), open book, on Day 2 from 4:30 pm.",
])

h1("Glossary")
B.append(("dl",[
 ("IoT","Internet of Things — the network of physical objects with sensors, software and connectivity that exchange data over the internet."),
 ("IIoT","Industrial IoT — IoT applied to industrial sectors such as manufacturing and energy."),
 ("Telemetry","The stream of metric readings a device sends to the cloud, as JSON key-value pairs."),
 ("Device token","The secret (dev_...) issued once at registration that authenticates every message a device sends."),
 ("MQTT","A lightweight publish/subscribe messaging protocol for constrained devices; a broker routes messages by topic."),
 ("REST API","An HTTP interface — POST /api/telemetry to write data, GET /api/device/state to read it."),
 ("Uplink / Downlink","Uplink = device → cloud telemetry. Downlink = cloud → device commands."),
 ("Virtual pin","A named key (V1, relay, pump) that dashboard control widgets write to and device code reacts to — Blynk-style."),
 ("Widget","A dashboard building block: number card, gauge, chart, LED, map (display) or button, switch, slider, terminal (control)."),
 ("Trigger","A rule that fires on an event — telemetry received, threshold crossed, device offline, command sent."),
 ("Alert rule","A trigger on a metric threshold or device-offline that tracks active alerts and can hand off to n8n."),
 ("n8n","A low-code workflow automation tool with 400+ integrations; flows start from a webhook and run nodes."),
 ("Webhook","A URL that starts an n8n workflow when the platform POSTs an event payload to it."),
 ("Workflow","A set of n8n nodes wired together that runs automatically on every trigger event."),
 ("AI node","A workflow node that calls a language model to summarise, classify or recommend actions from your data."),
 ("PWA","Progressive Web App — install the dashboard on your phone's home screen for mobile control."),
]))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    out.append(f"**IoT platform:** {C.PLATFORM_URL}  |  **Course page:** {C.COURSE_URL}")
    out.append("")
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd:
                    out+=["","   ```"]+["   "+ln for ln in cmd.split("\n")]+["   ```",""]
            out.append("")
        elif kind=="code": out+=["```",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="img": out+=[f"![{rest[1]}](courseware/assets/{rest[0]})","",f"*{rest[1]}*",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    out+=["---",f"*© 2026 {C.ORG}. All rights reserved. · www.tertiarycourses.com.sg*",""]
    return "\n".join(out)

MD_OUT=os.path.join(REPO,"LEARNER-GUIDE.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
    ("12","1 March 2025","Previous release — ThingSpeak/ESP8266/Node-RED based courseware.",C.TRAINER),
    ("13","6 July 2026",
     "Revamped to the IoTFlow platform (iot.tertiaryinfotech.com): 6 labs covering device registration, "
     "MQTT/HTTP telemetry, virtual-pin remote control, dashboards and n8n + AI automation.",C.TRAINER),
    (C.VERSION.lstrip("v"),C.VERSION_DATE,
     "Labs updated to the official IoTFlow Python client (pip install iotflow): send()/virtual_write() "
     "over HTTP, the telemetry_upload.py MQTT example, and @on_command control handlers; added the "
     "Python client setup section with the full example script.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph()
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            # explicit numbering so every lab's steps restart at 1 (List Number
            # style shares one numId and would run 1..44 across the document)
            para=doc.add_paragraph()
            para.paragraph_format.left_indent=DocxInches(0.32)
            para.paragraph_format.first_line_indent=DocxInches(-0.32)
            r=para.add_run(f"{i}."); r.bold=True
            para.add_run("  "+instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="img":
        path=os.path.join(ASSETS,rest[0])
        if os.path.exists(path):
            doc.add_picture(path,width=DocxInches(5.8))
            cap=doc.add_paragraph(); r=cap.add_run(rest[1]); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x55,0x5B,0x66)
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.FILE_STEM}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
